from fastapi import FastAPI, Depends, UploadFile, File, HTTPException, Form, Path
from sqlalchemy.orm import Session
from sqlalchemy import select, text
from contextlib import asynccontextmanager

# Fix the import - update this based on your actual file structure
from .db import get_db, engine, init_database  # Changed from .db to .database
from .models import Base, User, Document, ChatHistory
from .utils import hash_password, verify_password
from .embedding import chunk_text, get_embedding
from .mistral_api import ask_mistral
from .auth import create_access_token, get_current_user

import fitz
import pandas as pd
import os
import logging

# Set up logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

@asynccontextmanager
async def lifespan(app: FastAPI):
    """Handle startup and shutdown events"""
    # Startup
    logger.info("🚀 Starting application...")
    
    try:
        # Initialize database with pgvector extension
        init_database()
        logger.info("✅ Database initialized successfully")
    except Exception as e:
        logger.error(f"❌ Failed to initialize database: {e}")
        raise
    
    yield
    
    # Shutdown
    logger.info("🔄 Shutting down application...")

# Create FastAPI app with lifespan
app = FastAPI(
    title="Document Chatbot API",
    description="API for document-based chatbot with vector search",
    version="1.0.0",
    lifespan=lifespan
)

# Health check endpoint
@app.get("/health")
async def health_check():
    """Health check endpoint"""
    return {"status": "healthy", "message": "API is running"}

@app.post("/signup")
def signup(email: str = Form(...), password: str = Form(...), db: Session = Depends(get_db)):
    if db.query(User).filter(User.email == email).first():
        raise HTTPException(status_code=400, detail="Email already registered.")
    user = User(email=email, hashed_password=hash_password(password))
    db.add(user)
    db.commit()
    db.refresh(user)
    return {"message": "User created", "user_id": user.id}

@app.post("/login")
def login(email: str = Form(...), password: str = Form(...), db: Session = Depends(get_db)):
    user = db.query(User).filter(User.email == email).first()
    if not user or not verify_password(password, user.hashed_password):
        raise HTTPException(status_code=401, detail="Invalid credentials")
    token = create_access_token(data={"sub": str(user.id)})
    return {"access_token": token, "token_type": "bearer"}

@app.post("/ingest/")
def ingest_file(
    file: UploadFile = File(...),
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    temp_path = f"/tmp/{file.filename}"
    try:
        with open(temp_path, "wb") as f:
            f.write(file.file.read())

        ext = file.filename.lower().split(".")[-1]
        text = ""
        if ext == "txt":
            with open(temp_path, "r", encoding="utf-8") as f:
                text = f.read()
        elif ext == "pdf":
            doc = fitz.open(temp_path)
            for page in doc:
                text += page.get_text()
        elif ext == "csv":
            df = pd.read_csv(temp_path)
            text = df.astype(str).apply(" ".join, axis=1).str.cat(sep=" ")
        elif ext == "docx":
            from docx import Document as DocxDocument
            doc = DocxDocument(temp_path)
            text = "\n".join([para.text for para in doc.paragraphs])
        else:
            raise HTTPException(status_code=400, detail="Unsupported file type")

        chunks = chunk_text(text)
        for chunk in chunks:
            embedding = get_embedding(chunk).tolist()
            db.add(Document(chunk=chunk, embedding=embedding, source=file.filename, user_id=current_user.id))

        db.commit()
        return {"message": f"{file.filename} ingested and embeddings stored for user {current_user.id}."}
    finally:
        if os.path.exists(temp_path):
            os.remove(temp_path)

@app.get("/ask/")
def ask_question(query: str, db: Session = Depends(get_db), current_user: User = Depends(get_current_user)):
    query_vector = get_embedding(query).tolist()

    stmt = (
        select(Document)
        .filter(Document.user_id == current_user.id)
        .order_by(Document.embedding.l2_distance(query_vector))
        .limit(3)
    )
    top_docs = db.execute(stmt).scalars().all()

    if not top_docs:
        return {"answer": "No relevant information found in your documents."}

    top_chunks = [doc.chunk for doc in top_docs]
    answer = ask_mistral(query, top_chunks)

    # Save chat history
    chat = ChatHistory(question=query, answer=answer, user_id=current_user.id)
    db.add(chat)
    db.commit()

    return {"answer": answer}

@app.get("/history/")
def get_chat_history(db: Session = Depends(get_db), current_user: User = Depends(get_current_user)):
    history = db.query(ChatHistory).filter(ChatHistory.user_id == current_user.id).all()
    return [{"question": h.question, "answer": h.answer} for h in history]

@app.delete("/delete/{source}")
def delete_documents_by_source(
    source: str = Path(..., description="Source filename to delete all related documents"),
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    deleted_count = (
        db.query(Document)
        .filter(Document.user_id == current_user.id, Document.source == source)
        .delete()
    )
    db.commit()

    if deleted_count == 0:
        raise HTTPException(status_code=404, detail="No documents found for this source")

    return {"message": f"Deleted {deleted_count} documents for source '{source}'."}